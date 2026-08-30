from docx import Document
from docx.shared import Pt

out_file = r'c:\xampp\htdocs\pharmacy_management_prototype\Mergen_Pharmacy_System_Analysis_and_Design_Assignment.docx'

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title page
p = doc.add_paragraph()
p.alignment = 1
run = p.add_run('Mergen Pharmacy Management System\nSystem Analysis and Design Assignment')
run.bold = True
run.font.name = 'Calibri'
run.font.size = Pt(22)

doc.add_paragraph()

def add_label(text, size=12, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    return p

add_label('Prepared for: University / Course Requirement', 12, True)
add_label('Project Name: Pharmacy Management System Interface', 12)
add_label('Student / Group Name: __________________________', 12)
add_label('Submission Date: __________________________', 12)

doc.add_page_break()

# Section 1
add_label('1. Introduction', 16, True)
doc.add_paragraph('The Mergen Pharmacy project is a web-based pharmacy management system designed to support medicine catalog management, stock control, sales processing, supplier management, and reporting. The project demonstrates how a healthcare service can use digital tools to organize medication records, track low-stock and expiry situations, and simplify sales and purchase operations.')

doc.add_paragraph('The system is implemented using PHP, MySQL, HTML, CSS, JavaScript, and Chart.js, and is intended for local deployment through XAMPP. The interface allows users to move through key functions such as login, dashboard analytics, medicine creation, inventory monitoring, purchasing, sales, and report generation with a role-based access structure.')

# Section 2
add_label('2. Problem Statement', 16, True)
doc.add_paragraph('Pharmacies often manage medicine data, stock levels, sales, and suppliers using fragmented or paper-based processes. This causes issues such as delayed stock updates, poor visibility of expired products, inaccurate sales tracking, and difficulties in maintaining a consistent record of medicine movements.')

doc.add_paragraph('The system addresses these issues by providing a central digital platform to organize medicine records, visualize inventory conditions, automate sales workflows, and support reporting for operational decision-making.')

# Section 3
add_label('3. Objectives of the System', 16, True)
for item in [
    'To provide a clear and responsive interface for pharmacy operations.',
    'To manage the medicine catalog with details such as quantity, price, expiry date, and category.',
    'To monitor inventory status including in-stock, low-stock, out-of-stock, and expired products.',
    'To process point-of-sale transactions efficiently through a simplified sales workflow.',
    'To maintain supplier and purchase information for smooth stock replenishment.',
    'To generate analytical reports for sales, inventory, purchases, and expiry tracking.',
    'To support role-based access and tailored user views for administrators, pharmacists, and staff.'
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 4
add_label('4. Scope of the System', 16, True)
doc.add_paragraph('The project covers the design and interface prototype of a pharmacy management system. It includes login functionality, dashboard analytics, medicine and inventory management, supplier and purchase handling, sales workflow, and reporting modules. It is suitable for academic demonstration, business presentation, and system analysis practice.')

doc.add_paragraph('The system does not replace a full production-grade ERP but instead demonstrates the architecture and operational flow of a pharmacy digital platform.')

# Section 5
add_label('5. Functional Requirements', 16, True)
for item in [
    'Admin, pharmacist, and staff users must be able to log in securely with role-based access.',
    'The system must allow the addition, editing, and viewing of medicine information.',
    'The system must support stock adjustment and display stock status alerts.',
    'The system must allow sales entry through a POS-style interface including product search and quantity update.',
    'The system must record purchase orders and supplier details.',
    'The system must provide reports based on sales, purchases, inventory, low-stock, and expiry.',
    'The system must allow role switching in the interface to demonstrate different access levels.'
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 6
add_label('6. Non-Functional Requirements', 16, True)
for item in [
    'The interface should be user-friendly, clean, and easy to navigate.',
    'The application should render quickly on a local web server using PHP and MySQL.',
    'The design should be responsive and consistent across desktop displays.',
    'The product should support secure session handling and role-based menu visibility.',
    'The system should be easy to maintain and extend for future production deployment.'
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 7
add_label('7. System Architecture and Design Overview', 16, True)
doc.add_paragraph('The proposed system consists of a web-based front end, a PHP business layer, and a MySQL database. The front end is responsible for user interaction, while the database stores medicine, sale, purchase, supplier, and staff records. The architecture is simple yet scalable and suitable for classroom demonstration and prototype deployment.')

doc.add_paragraph('Diagram Placeholder: High-Level Architecture Diagram')
placeholder = doc.add_table(rows=1, cols=1)
placeholder.style = 'Table Grid'
placeholder.rows[0].cells[0].text = 'Insert System Architecture Diagram here\n\n[User Interface / Browser]\n\nPHP Application Layer\n\nMySQL Database\n\nReports, POS, Inventory, Authentication'

# Section 8
add_label('8. Use-Case Description', 16, True)
doc.add_paragraph('The system supports several primary user roles: admin, pharmacist, and staff. An admin can manage medicines, suppliers, staff, and reports. A pharmacist can search medicines, check inventory, and generate sales. A staff user can assist with counter work and stock visibility tasks.')

doc.add_paragraph('Typical use cases include logging in, creating a medicine record, checking low stock, adding a supplier, processing a sale, viewing a report, and switching roles to test permission controls.')

doc.add_paragraph('Diagram Placeholder: Use-Case Diagram')
usecase = doc.add_table(rows=1, cols=1)
usecase.style = 'Table Grid'
usecase.rows[0].cells[0].text = 'Insert Use-Case Diagram here\n\nActors: Admin, Pharmacist, Staff\nUse Cases: Login, Manage Medicines, Process Sale, Manage Inventory, View Reports, Manage Suppliers'

# Section 9
add_label('9. Module Breakdown', 16, True)
for item in [
    'Authentication Module: Login, session management, quick demo logins, and password reset interface.',
    'Dashboard Module: KPI cards, chart visualizations, and summary analytics.',
    'Medicine Management Module: Add, edit, delete, search, and medicine profile details.',
    'Inventory Module: Stock level monitoring, expiry management, and status badges.',
    'Sales Module: POS sale workflow, cart handling, payment method, and thermal receipt generation.',
    'Purchase and Supplier Module: Purchase ordering, supplier directory, and stock replenishment records.',
    'Reports Module: Sales, purchases, inventory, low stock, and expiry report tabs.',
    'Profile and User Settings Module: Profile management and access customization.'
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 10
add_label('10. Database Design and Data Flow', 16, True)
doc.add_paragraph('The database is designed using normalized tables to manage users, medicines, purchases, sales, suppliers, and inventory status. Core entities include users, medicines, suppliers, purchases, sales, and staff records. These are linked so that stock adjustments and transactions can be tracked over time.')

doc.add_paragraph('Diagram Placeholder: ER Diagram / Data Flow Diagram')
flow = doc.add_table(rows=1, cols=1)
flow.style = 'Table Grid'
flow.rows[0].cells[0].text = 'Insert ER Diagram or DFD here\n\nEntities: User, Medicine, Supplier, Purchase, Sale, Inventory, Report\n\nUser -> Login -> Dashboard -> Sales/Inventory/Reports\nMedicine -> Stock -> Sales/Purchases\nSupplier -> Purchase -> Inventory'

# Section 11
add_label('11. User Interface Evaluation', 16, True)
doc.add_paragraph('The interface emphasizes clarity, healthcare-related styling, and operational convenience. The color palette uses teal and emerald tones to maintain a professional medical context. Navigation is structured with a sidebar and top navbar, ensuring users can quickly reach essential modules without confusion.')

doc.add_paragraph('The system also includes a role switcher and quick demo authentication for academic demonstrations, which helps instructors evaluate the application with minimal setup.')

# Section 12
add_label('12. Security and Maintenance Considerations', 16, True)
for item in [
    'Use role-based access control to restrict modules by user type.',
    'Validate and sanitize all user input before database operations.',
    'Store passwords using secure hashing in production versions.',
    'Add logging for transactions and critical system events.',
    'Continue database backups and audit checks for long-term reliability.',
    'Implement encryption and session protection in future production releases.'
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 13
add_label('13. Proposed Future Enhancements', 16, True)
for item in [
    'Barcode scanner integration for rapid medicine lookup and sales entry.',
    'Online payment support such as SSLCommerz, bKash, or card gateways.',
    'Automated alerts for low stock and medicine expiry.',
    'Enhanced security with password hashing and audit logs.',
    'Advanced analytics dashboard and business intelligence panels.'
]:
    doc.add_paragraph(item, style='List Bullet')

# Conclusion
add_label('14. Conclusion', 16, True)
doc.add_paragraph('The Mergen Pharmacy system demonstrates a practical and visually polished pharmacy management interface that covers essential business flow in a healthcare setting. The project successfully combines data management, user interaction, reporting, and role-driven navigation in a single prototype environment.')

doc.add_paragraph('This assignment highlights the system analysis and design thinking behind the project and provides a good foundation for further development into a complete, production-ready pharmacy management application.')

# References
add_label('15. References', 16, True)
doc.add_paragraph('Project README: Mergen Pharmacy Management System Interface')
doc.add_paragraph('Project Source Files: dashboard.php, login.php, medicines/, sales/, inventory/, suppliers/, reports/, includes/, config/, database/')
doc.add_paragraph('XAMPP phpMyAdmin setup and PHP/MySQL local deployment workflow')

doc.save(out_file)
print(f'Created assignment document: {out_file}')
